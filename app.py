from flask import Flask, render_template, request, send_file
import fitz  # PyMuPDF
from docx import Document
import os
import pdfplumber
import pandas as pd
from pdf2image import convert_from_path
import pytesseract

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def make_columns_unique(cols):
    """Ensure all column names are unique by adding suffixes to duplicates."""
    seen = {}
    new_cols = []
    for c in cols:
        if c in seen:
            seen[c] += 1
            new_cols.append(f"{c}_{seen[c]}")
        else:
            seen[c] = 0
            new_cols.append(c)
    return new_cols


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'pdf_file' not in request.files:
            return "No file uploaded"

        pdf_file = request.files['pdf_file']
        if pdf_file.filename == '':
            return "No selected file"

        pdf_path = os.path.join(UPLOAD_FOLDER, pdf_file.filename)
        pdf_file.save(pdf_path)

        conversion_type = request.form.get("conversion_type")

        # -------- PDF to Word --------
        if conversion_type == "word":
            docx_filename = pdf_file.filename.rsplit('.', 1)[0] + '.docx'
            docx_path = os.path.join(UPLOAD_FOLDER, docx_filename)

            doc = Document()
            with fitz.open(pdf_path) as pdf:
                for page in pdf:
                    text = page.get_text()
                    if text.strip():
                        doc.add_paragraph(text)

            doc.save(docx_path)
            return send_file(docx_path, as_attachment=True)

        # -------- PDF to Excel --------
        elif conversion_type == "excel":
            excel_filename = pdf_file.filename.rsplit('.', 1)[0] + '.xlsx'
            excel_path = os.path.join(UPLOAD_FOLDER, excel_filename)
            all_tables = []

            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for t in tables:
                        if len(t) < 2:
                            continue  # skip empty or single-row tables
                        df = pd.DataFrame(t[1:], columns=make_columns_unique(t[0]))
                        all_tables.append(df)

            if all_tables:
                final_df = pd.concat(all_tables, ignore_index=True)
                final_df.to_excel(excel_path, index=False)
            else:
                pd.DataFrame().to_excel(excel_path, index=False)

            return send_file(excel_path, as_attachment=True)

        # -------- PDF OCR to Word --------
        elif conversion_type == "ocr":
            docx_filename = pdf_file.filename.rsplit('.', 1)[0] + '_ocr.docx'
            docx_path = os.path.join(UPLOAD_FOLDER, docx_filename)

            images = convert_from_path(pdf_path)
            doc = Document()

            for img in images:
                text = pytesseract.image_to_string(img)
                if text.strip():
                    doc.add_paragraph(text)

            doc.save(docx_path)
            return send_file(docx_path, as_attachment=True)

        else:
            return "Invalid conversion type selected"

    return render_template('index.html')


if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True, port=8000)

